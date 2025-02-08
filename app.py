from docx.shared import Inches
from flask import Flask, render_template, request
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import smtplib
import os
import datetime
import re

app = Flask(__name__, template_folder='templates')

def is_valid_email(email):
    pattern = r'^[\w\.-]+@[a-zA-Z\d\.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_and_send_letter', methods=['POST'])
def generate_and_send_letter():
    required_fields = ['complaint_type', 'complainant_name', 'complainant_email',
                       'from_address', 'recipient_email', 'additional_info']

    if not all(field in request.form for field in required_fields):
        return "Incomplete form data"

    complaint_type = request.form['complaint_type']
    complainant_name = request.form['complainant_name']
    complainant_email = request.form['complainant_email']
    from_address = request.form['from_address']
    recipient_email = request.form['recipient_email']
    additional_info = request.form['additional_info']

    if not is_valid_email(recipient_email):
        return render_template('error.html', error="Invalid recipient email address")

    if complaint_type == 'financial':
        letter_content, logo_path = generate_financial_letter(complainant_name, complainant_email, from_address, additional_info)
    elif complaint_type == 'non_financial':
        letter_content, logo_path = generate_non_financial_letter(complainant_name, complainant_email, from_address, additional_info)
    else:
        return "Invalid complaint type"

    filename = save_letter_to_file(letter_content, complainant_name, complaint_type, logo_path)

    sender_email = "saminathan.mxiib@gmail.com"  # Update with your email
    sender_password = "eycj edes gmlc rkka"  # Update with your app password

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = "Complaint Letter"

    body = f"Dear recipient,\n\nPlease find attached the complaint letter and take action immediately.\n\nThank you,\nYours faithfully,\nTamil Nadu Cyber Crime"
    msg.attach(MIMEText(body, 'plain'))

    with open(filename, 'rb') as f:
        attachment = MIMEApplication(f.read(), _subtype='docx')
        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
        msg.attach(attachment)

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(sender_email, sender_password)
            smtp.send_message(msg)

        return render_template('letter_generated.html', message="Letter generated and sent",
                               recipient_email=recipient_email)
    except Exception as e:
        return render_template('error.html', error=f"Error sending email: {str(e)}")

def generate_financial_letter(complainant_name, complainant_email, from_address, additional_info):
    logo_path = r'C:\Users\HP\Desktop\Practice\PycharmProjects\pythonProject\logo_crime.jpeg'  # Update with your logo path
    return f"Dear Bank,\n\nWe have received a complaint from {complainant_name} ({complainant_email}) regarding a financial fraud. Details: {additional_info}\n\nThank you,\nYours faithfully,\nTamil Nadu Cyber Crime", logo_path

def generate_non_financial_letter(complainant_name, complainant_email, from_address, additional_info):
    logo_path = r'C:\Users\HP\Desktop\Practice\PycharmProjects\pythonProject\logo_crime.jpeg'  # Update with your logo path
    return f"Dear Social Media Company,\n\nWe have received a complaint from {complainant_name} ({complainant_email}) regarding non-financial fraud. Details: {additional_info}\n\nThank you,\nYours faithfully,\nTamil Nadu Cyber Crime", logo_path

def save_letter_to_file(letter_content, complainant_name, complaint_type, logo_path):
    try:
        date = datetime.datetime.now()
        doc = Document()

        doc.add_picture(logo_path, width=Inches(1.0))
        doc.add_heading('Complaint Letter from TamilNadu Cyber Crime', level=0)
        doc.add_paragraph(date.strftime("%B %d, %Y"))
        doc.add_paragraph("\n" + letter_content)

        filename = f"{complaint_type.capitalize()}_{complainant_name}_{date.strftime('%Y%m%d')}.docx"
        save_path = r'C:\Users\HP\Desktop\Practice\PycharmProjects\pythonProject\letters'  # Update with your desired path
        absolute_path = os.path.join(save_path, filename)

        os.makedirs(save_path, exist_ok=True)

        # Print the absolute path for debugging purposes
        print("Absolute Path:", absolute_path)

        doc.save(absolute_path)
        return absolute_path  # Return the absolute path of the saved file
    except FileNotFoundError as fnf_error:
        return f"FileNotFoundError: {fnf_error}. Make sure the directory exists and check the path."
    except Exception as e:
        return f"Error saving letter: {str(e)}"



if __name__ == '__main__':
    app.run(debug=True)
